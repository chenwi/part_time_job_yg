# -*- encoding: utf-8 -*-
import pandas as pd
from bs4 import BeautifulSoup
import re
import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
import win32com
from win32com.client import DispatchEx, constants
from tkinter import Tk, Label, Button, StringVar, Entry
from tkinter.filedialog import askdirectory
import threading
import time
import pythoncom

win32com.client.gencache.EnsureDispatch('Word.Application')


def sub(s):
    s = re.sub(r'\xa0', ' ', s)
    if s == 'nan':
        s = ''
    return s


class Extractor:
    def __init__(self, path):
        self.path = path

    def read_html(self):
        """读取html,解析html导入到dataframe，最后转化为list"""
        with open(self.path, 'r') as f:
            html = f.read()
        td = pd.read_html(html)[0]
        td = td.applymap(str)
        td = td.values.tolist()
        return td

    def used_data(self):
        """used data, used columns, 筛选有用的列"""
        pass

    def get_table(self):
        """clean data, 数据转化为table，type：list"""
        pass


class AbsQuant(Extractor):
    columns = ['Position', 'Sample Name', 'Gene Name', 'Cq', 'Concentration', 'Call',  # 'Excluded',
               'Sample Type', 'Standard', 'Cq Mean', 'Cq Error', 'Concentration Mean', 'Concentration Error',
               'Replicate Group', 'Dye',  # 'Edited Call',
               'Slope', 'EPF', 'Failure', 'Notes', 'Prep Notes', 'Number']

    def __init__(self, path):
        super(AbsQuant, self).__init__(path)
        self.path = path
        self.all_table = self.get_all_table()
        self.table = self.get_table()
        self.img_src = os.path.join(os.path.abspath(os.path.dirname(self.path)), self.get_img())

    def used_data(self):
        td = self.read_html()
        alldata = []
        flag = False
        for row in td:
            if row[2] == 'Color':
                flag = True
                continue
            if flag:
                if row[2].startswith('Statistical'):
                    break
                items = row[5:11] + row[17:25] + row[31:37]
                alldata.append(items)
        return alldata

    def get_img(self):
        with open(self.path, 'r') as f:
            html = f.read()
        soup = BeautifulSoup(html, 'lxml')
        tables = soup.findAll('table')
        tab = tables[0]
        flag = False
        for tr in tab.findAll('tr'):
            if flag:
                img = tr.findAll('img')
                src = img[0].get('src')
                return src
            for td in tr.findAll('td'):
                text = td.getText()
                if text.startswith('Amplification'):
                    flag = True
                    break

    def get_all_table(self):
        alldata = self.used_data()
        all_table = [self.columns]
        for row in alldata:
            row = list(map(sub, row))
            all_table.append(row)
        return all_table

    def get_table(self):
        table = []
        for row in self.all_table:
            items = row[:6] + row[13:14]
            table.append(items)
        return table

    def all_table2df(self):
        df = pd.DataFrame(self.all_table[1:], columns=self.columns)
        return df

    def to_csv(self, path='all.csv'):
        df = self.all_table2df()
        df.to_csv(path, index=None)

    # def html2pdf(self):
    #     df = self.all_table2df()
    #     a = df.to_html(index=False)
    #     # print(a)
    #     import pdfkit
    #     options = {
    #         'page-size': 'Letter',
    #         'encoding': "UTF-8",
    #         'custom-header': [
    #             ('Accept-Encoding', 'gzip')
    #         ]
    #     }
    #     path_wk = r'D:\wkhtmltox\bin\wkhtmltopdf.exe'  # 安装位置
    #     config = pdfkit.configuration(wkhtmltopdf=path_wk)
    #     pdfkit.from_string(a, "test.pdf", configuration=config, options=options)


class BasicInfo(Extractor):
    def __init__(self, path):
        super(BasicInfo, self).__init__(path)
        self.table = self.get_table()

    def used_data(self):
        td = self.read_html()
        used_data = td[1:2] + td[5:7] + td[8:10]
        return used_data

    def get_table(self):
        td = self.used_data()
        table = []
        for row in td:
            temp = []
            for val in row:
                if val != 'nan':
                    val = sub(val)
                    temp.append(val)
            table.append(temp)
        return table


class RunEditor(Extractor):
    def __init__(self, path):
        super(RunEditor, self).__init__(path)
        self.table = self.get_table()

    def used_data(self):
        td = self.read_html()
        flag = False
        used_data = []
        for row in td:
            if row[2] == 'Programs':
                flag = True
                continue
            if flag:
                if row[2].startswith('Temperature'):
                    break
                used_data.append(row[2:6])

        return used_data

    def table_filter(self, table):
        oktab = []
        for row in table:
            row = list(map(sub, row))
            oktab.append(row[1:])
        return oktab

    def get_table(self):
        td = self.used_data()
        tables = []
        i = 0
        while i < len(td):
            if td[i][0] != 'nan':
                table = []
                title = sub(td[i][0])
                i += 1
                while i < len(td) and td[i][0] == 'nan':
                    table.append(td[i])
                    i += 1
                tables.append((title, self.table_filter(table)))

        return tables


class Writer:
    def __init__(self, file_dir, save_dir):
        """
        :word文件转pdf
        :param doc_name word文件名称
        :param pdf_name 转换后pdf文件名称
        """
        name = os.path.split(file_dir)[-1]
        doc_name = name + '.docx'
        pdf_name = name + '.pdf'

        self.file_dir = os.path.abspath(file_dir)
        self.save_dir = os.path.abspath(save_dir)
        self.doc_path = os.path.join(self.save_dir, doc_name)
        self.pdf_path = os.path.join(self.save_dir, pdf_name)

    def write_doc(self):
        # 转docx
        inpath = os.path.join(self.file_dir, 'report_resources')
        basic_path = os.path.join(inpath, 'basic_info.html')
        run_path = os.path.join(inpath, 'run_editor.html')
        abs_path = os.path.join(inpath, 'abs quant001.html')
        basic_info = BasicInfo(basic_path)
        run_editor = RunEditor(run_path)
        abs_quant = AbsQuant(abs_path)

        basic_table = basic_info.table
        run_table = run_editor.table
        abs_table = abs_quant.table
        abs_img_src = abs_quant.img_src
        ###all csv
        # abs_quant.to_csv(os.path.join(self.file_dir,os.path.split(self.file_dir)[-1]+'.csv'))

        doc = docx.Document()  # 新建文档
        doc.styles['Normal'].font.name = u'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        p = doc.add_paragraph('')
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run('荧光定量 PCR 检测报告')
        r.font.size = Pt(14)  # 设置字号
        r.font.bold = True  # 加粗

        doc.add_paragraph('')
        ### part 1
        p = doc.add_paragraph('')
        r = p.add_run(r'1.   Info（实验信息）')
        r.font.bold = True  # 加粗

        basic_tab_rows, basic_tab_cols = len(basic_table), len(basic_table[0])
        basic_tab = doc.add_table(rows=basic_tab_rows, cols=basic_tab_cols, style='Table Grid')

        for r in range(basic_tab_rows):
            for c in range(basic_tab_cols):
                basic_tab.cell(r, c).text = basic_table[r][c]
        doc.add_paragraph('')  # blank line

        # part 2
        p = doc.add_paragraph('')
        r = p.add_run(r'2. Run Profile（程序设置）')
        r.font.bold = True  # 加粗
        p = doc.add_paragraph('')
        r = p.add_run(r'Programs')
        r.font.bold = True
        for (title, td) in run_table:
            rows, cols = len(td), len(td[0])
            p = doc.add_paragraph('')
            p.add_run(title)
            table = doc.add_table(rows=rows, cols=cols, style='Table Grid')

            for r in range(rows):
                for c in range(cols):
                    table.cell(r, c).text = td[r][c]
            doc.add_paragraph('')
        # part 3
        doc.add_page_break()
        p = doc.add_paragraph('')
        r = p.add_run(r'3. Analysis（实验分析）')
        r.font.bold = True  # 加粗
        p = doc.add_paragraph('')
        r = p.add_run(r'Abs Quant（定量分析）')
        r.font.bold = True  # 加粗
        # Curves
        p = doc.add_paragraph('')
        p.add_run(r'Ampification Curves（扩增曲线）')

        p = doc.add_picture(abs_img_src, width=Inches(6))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Result Table
        p = doc.add_paragraph('')
        p.add_run(r'Result Table（实验结果）')

        abs_tab_rows, abs_tab_cols = len(abs_table), len(abs_table[0])
        abs_tab = doc.add_table(rows=abs_tab_rows, cols=abs_tab_cols, style='Table Grid')

        for r in range(abs_tab_rows):
            for c in range(abs_tab_cols):
                run = abs_tab.cell(r, c).paragraphs[0].add_run(abs_table[r][c])
                if r == 0:
                    run.font.bold = 1
                    abs_tab.cell(r, c).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        abs_tab.cell(0, 1).width = Cm(5)
        abs_tab.cell(0, 2).width = Cm(4)
        doc.save(self.doc_path)

    def write_pdf(self):
        pythoncom.CoInitialize()
        word = DispatchEx("Word.Application")
        pythoncom.CoInitialize()
        try:
            if os.path.exists(self.pdf_path):
                os.remove(self.pdf_path)
            worddoc = word.Documents.Open(self.doc_path, ReadOnly=1)
            worddoc.SaveAs(self.pdf_path, FileFormat=17)
            worddoc.Close()
            # return pdf_name
        except Exception as e:
            # print('aaa')
            print(e)
        finally:
            word.Quit(constants.wdDoNotSaveChanges)


class Gui:

    def __init__(self, root):
        self.root = root
        self.path = StringVar()
        self.outpath = StringVar()
        self.file_dir = ''
        self.save_dir = ''
        Label(self.root, text="").grid(row=0, column=0)
        Label(self.root, text="原始文件:").grid(row=1, column=0)
        Entry(self.root, textvariable=self.path, width=40).grid(row=1, column=1)
        Button(self.root, text="打开", command=self.selectPath).grid(row=1, column=3)
        Label(self.root, text="保存文件:").grid(row=2, column=0)
        Entry(self.root, textvariable=self.outpath, width=40).grid(row=2, column=1)
        Button(self.root, text="保存", command=self.select_outPath).grid(row=2, column=3)
        self.lb = Label(self.root, text='欢迎 !')
        self.lb.grid(row=3, column=1)
        self.bt = Button(self.root, text="生成报告", command=self.run)
        self.bt.grid(row=4, column=1)
        self.write_ok = False
        self.err = False

    def selectPath(self):
        self.file_dir = askdirectory()
        self.save_dir = self.file_dir
        self.path.set(self.file_dir)
        self.outpath.set(self.file_dir)
        self.err = False

    def select_outPath(self):
        self.save_dir = askdirectory()
        self.outpath.set(self.save_dir)

    def run(self):
        self.lb.config(text="运行...")
        th = threading.Thread(target=self.write)
        th.start()
        th2 = threading.Thread(target=self.ui)
        th2.start()

    def ui(self):
        if not self.write_ok:
            self.bt.config(state='disabled')
        while not self.write_ok:
            time.sleep(3)
        self.bt.config(state='normal')
        if not self.err:
            self.lb.config(text="完成！")

    def write(self):
        self.write_ok = False
        try:
            writer = Writer(self.file_dir, self.save_dir)
            writer.write_doc()
            writer.write_pdf()
            self.write_ok = True

        except Exception as e:
            print(e)
            self.err = True
            self.lb.config(text="路径或文件错误！")
            self.write_ok = True


if __name__ == '__main__':
    root = Tk()
    root.title('LC96 报告生成工具')
    root.geometry("400x150")
    root.resizable(width=True, height=False)
    Gui(root)
    root.mainloop()
