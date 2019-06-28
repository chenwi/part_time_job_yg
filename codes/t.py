from tkinter import Tk, Label, Button, StringVar, Entry
from tkinter.filedialog import askdirectory
import threading
import time
import docx


class Gui:

    def __init__(self, root):
        self.root = root
        self.path = StringVar()
        self.outpath = StringVar()
        self.file_dir = ''
        self.save_dir = ''
        Label(self.root, text="").grid(row=0, column=0)
        Label(self.root, text="原始文件:").grid(row=1, column=0)
        Entry(self.root, textvariable=self.path, width=40).grid(row=1, column=1)
        Button(self.root, text="打开", command=self.selectPath).grid(row=1, column=3)
        Label(self.root, text="保存文件:").grid(row=2, column=0)
        Entry(self.root, textvariable=self.outpath, width=40).grid(row=2, column=1)
        Button(self.root, text="保存", command=self.select_outPath).grid(row=2, column=3)
        self.lb = Label(self.root, text='欢迎 !')
        self.lb.grid(row=3, column=1)
        self.bt = Button(self.root, text="生成报告", command=self.run)
        self.bt.grid(row=4, column=1)
        self.write_ok = False
        self.err = False

    def selectPath(self):
        self.file_dir = askdirectory()
        self.save_dir = self.file_dir
        self.path.set(self.file_dir)
        self.outpath.set(self.file_dir)
        self.err = False
        # print(self.path)
        # print(self.file_dir)
        # print(self.save_dir)

    def select_outPath(self):
        self.save_dir = askdirectory()
        self.outpath.set(self.save_dir)

    def run(self):
        self.lb.config(text="运行...")
        th = threading.Thread(target=self.write)
        th2 = threading.Thread(target=self.ui)
        # th2.setDaemon(True)
        th.start()
        th2.start()
        # for t in [th,th2]:
        #     t.join()

    def ui(self):
        if not self.write_ok:
            self.bt.config(state='disabled')
            # print(0)
        while not self.write_ok:
            time.sleep(6)
            print(1)
        self.bt.config(state='normal')
        if not self.err:
            self.lb.config(text="完成！")

    def write(self):
        self.write_ok = False
        try:
            # self.lb.config(text="running...")
            doc = docx.Document()  # 新建文档
            doc.styles['Normal'].font.name = u'Times New Roman'
            # doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            p = doc.add_paragraph('')
            # p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p.add_run('荧光定量 PCR 检测报告')
            print(self.file_dir)
            doc.save(self.file_dir+r'/1.docx')
            # self.lb.config(text="finished")
            self.write_ok = True
            # print('ok')
        except Exception as e:
            print(e)
            self.err = True
            self.lb.config(text="路径或文件错误！")
            self.write_ok = True


if __name__ == '__main__':
    root = Tk()
    root.title('LC96 报告生成工具')
    root.geometry("400x150")
    Gui(root)
    root.mainloop()
