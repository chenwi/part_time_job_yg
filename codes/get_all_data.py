from helper import *
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt,Inches,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import os


def __all_data(path):
    # print(html)
    td = read_html(path)
    # td = pd.read_html(html)[0]
    # print(td.head)
    # print(td[4:5])
    # print(td[16:17])
    # a=pd.DataFrame(td[0:20])
    # a=a.values.tolist()
    # print(a)
    # for i in a:
    #     print(i)
    alldata = []
    flag = False
    for row in td:
        # print(row)
        if row[2] == 'Color':
            flag = True
            continue
        if flag:
            if row[2].startswith('Statistical'):
                # print('\n11111\n',row)
                break
            # print(row)
            items = row[5:11] + row[17:25] + row[31:37]
            alldata.append(items)
    return alldata


def get_img(path):
    with open(path, 'r') as f:
        html = f.read()
    soup = BeautifulSoup(html, 'lxml')
    tables = soup.findAll('table')
    tab = tables[0]
    flag = False
    for tr in tab.findAll('tr'):
        if flag:
            img = tr.findAll('img')
            src=img[0].get('src')
            return src
        for td in tr.findAll('td'):
            text = td.getText()
            if text.startswith('Amplification'):
                # print(text)
                flag = True
                break


def get_all_table(path):
    alldata = __all_data(path=path)
    table = []
    for row in alldata:
        row = list(map(sub, row))
        # print(row)
        table.append(row)
    return table


# tables = get_all_tables()
columns = ['Position', 'Sample Name', 'Gene Name', 'Cq', 'Concentration', 'Call',  # 'Excluded',
           'Sample Type', 'Standard', 'Cq Mean', 'Cq Error', 'Concentration Mean', 'Concentration Error',
           'Replicate Group', 'Dye',  # 'Edited Call',
           'Slope', 'EPF', 'Failure', 'Notes', 'Prep Notes', 'Number']


def all_table2df(table):
    df = pd.DataFrame(table, columns=columns)
    return df


def to_csv(table, path=None):
    df = all_table2df(table)
    # import numpy as np
    # df = df.replace('nan', '')
    # df['Number'] = df['Number'].astype(float)
    # df['Number'] = df['Number'].astype(int)
    # print(df.head())
    df.to_csv('a.csv', index=None)


def html2pdf(table):
    df = all_table2df(table)
    a = df.to_html(index=False)
    # print(a)
    import pdfkit
    options = {
        'page-size': 'Letter',
        'encoding': "UTF-8",
        'custom-header': [
            ('Accept-Encoding', 'gzip')
        ]
    }
    path_wk = r'D:\wkhtmltox\bin\wkhtmltopdf.exe'  # 安装位置
    config = pdfkit.configuration(wkhtmltopdf=path_wk)
    pdfkit.from_string(a, "test.pdf", configuration=config, options=options)


def all_table2doc(table, img, path='all_data.docx'):
    # 转docx
    doc = docx.Document()  # 新建文档
    p = doc.add_paragraph('')
    r = p.add_run(r'3. Analysis（实验分析）')
    r.font.bold = True  # 加粗
    p = doc.add_paragraph('')
    r = p.add_run(r'Abs Quant')
    r.font.bold = True  # 加粗
    #Curves
    p = doc.add_paragraph('')
    p.add_run(r'Ampification Curves')
    p=doc.add_picture(img,width=Inches(6))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Result Table
    p = doc.add_paragraph('')
    p.add_run(r'Result Table')

    td = []
    table.insert(0, columns)
    for row in table:
        # items = row[:6] + row[13:14]
        items=row
        td.append(items)

    rows, cols = len(td), len(td[0])
    # print(rows, cols)
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
    table.columns[1].width = Pt(2)
    table.columns[2].width = Pt(2)
    # 遍历表格
    for r in range(rows):
        for c in range(cols):
            run = table.cell(r, c).paragraphs[0].add_run(td[r][c])  # 填入的内容
            if r == 0:
                # run.font.name = 'Times New Roman'  # 设置字体
                # run.font.size = Pt(11)  # 设置字号
                run.font.bold = 1  # 设置是否加粗
                # run.font.color.rgb = RGBColor.from_string('00BFFF')  # 设置文字颜色
                table.cell(r, c).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置居中
    # for r, row in enumerate(table.rows):
    #     for c in range(len(row.cells)):
    #         # print(c)
    #         row.cells[c].text = td[r][c]
    table.cell(0, 1).width = Cm(5)
    table.cell(0, 2).width = Cm(4)
    # table.autofit = True
    doc.save(path)



if __name__ == '__main__':
    path = r'data/Demo_10-Fold Dilution/report_resources/abs quant001.html'
    # path = r'data/Demo_Qual Detect Mono Color/report_resources/abs quant001.html'
    table = get_all_table(path=path)
    for i in table:
        print(i)
    # print(tables)
    src = get_img(path)
    print(src)
    img_path = os.path.join(os.path.dirname(path), src)
    all_table2doc(table,img_path)
    to_csv(table)

    # with open(path, 'r') as f:
    #     html = f.read()
    # soup = BeautifulSoup(html, 'lxml')
    # # img = soup.find_all('img')
    # # src=img[0].get('src')
    # # for i in range(len(img)):
    # #     print(img[i].get('src'))
    # tables = soup.findAll('table')
    # tab = tables[0]
    # flag = False
    # for tr in tab.findAll('tr'):
    #     if flag:
    #         img = tr.findAll('img')
    #         print(tr)
    #         print(tr.get('src'))
    #         print(img[0].get('src'))
    #         break
    #     for td in tr.findAll('td'):
    #         text = td.getText()
    #         if text.startswith('Amplification'):
    #             print(text)
    #             flag = True
    #             break
