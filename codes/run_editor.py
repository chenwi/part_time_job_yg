from helper import *
import docx
from bs4 import BeautifulSoup

def __run_used_data(path):
    td = read_html(path)
    flag = False
    used_data = []
    for row in td:
        if row[2] == 'Programs':
            flag = True
            continue
        if flag:
            if row[2].startswith('Temperature'):
                break
            used_data.append(row[2:6])

    return used_data

def table_filter(table):
    oktab = []
    for row in table:
        row = list(map(sub, row))
        oktab.append(row[1:])
    return oktab


def get_basic_table(path):
    td = __run_used_data(path)
    tables = []
    i = 0
    while i < len(td):
        if td[i][0] != 'nan':
            table = []
            title = sub(td[i][0])
            i += 1
            while i < len(td) and td[i][0] == 'nan':
                table.append(td[i])
                i += 1
            tables.append((title, table_filter(table)))

    return tables


def basic_table2doc(tables, path='run_data.docx'):
    # 转docx
    doc = docx.Document()  # 新建文档
    p = doc.add_paragraph('')
    r = p.add_run(r'2. Run Profile（程序设置）')
    r.font.bold = True  # 加粗
    p = doc.add_paragraph('')
    r = p.add_run(r'Programs')
    r.font.bold = True  # 加粗
    for (title,td) in tables:
        rows, cols = len(td), len(td[0])
        print(rows, cols)
        p = doc.add_paragraph('')
        p.add_run(title)
        # r.font.bold = True  # 加粗
        table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
        # 遍历表格
        for r in range(rows):
            for c in range(cols):
                table.cell(r, c).text = td[r][c]
        doc.add_paragraph('')
    doc.save(path)


if __name__ == '__main__':
    # path = r'data/Demo_Qual Detect Mono Color/report_resources/run_editor.html'
    path=r'data/Demo_2-Fold Dilution/report_resources/run_editor.html'
    a = __run_used_data(path)
    for i in a:
        print(i)
    b = get_basic_table(path)
    basic_table2doc(b)

