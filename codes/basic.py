from helper import *
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor


def __basic_used_data(path):
    td = read_html(path)
    used_data = td[1:2] + td[5:7] + td[8:10]
    return used_data


def get_basic_table(path):
    td = __basic_used_data(path)
    table = []
    for row in td:
        temp = []
        for val in row:
            if val != 'nan':
                val = sub(val)
                temp.append(val)
        table.append(temp)
    return table

def basic_table2doc(td, path='basic_data.docx'):
    # 转docx
    doc = docx.Document()  # 新建文档
    rows, cols = len(td), len(td[0])
    print(rows, cols)
    title = r'1.   Info（实验信息）'
    p = doc.add_paragraph('')
    r = p.add_run(title)
    r.font.bold = True  # 加粗

    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
    # 遍历表格
    for r in range(rows):
        for c in range(cols):
            table.cell(r, c).text=td[r][c]
            # run = table.cell(r, c).paragraphs[0].add_run(td[r][c])  # 填入的内容
            # if r == 0:
            #     # run.font.name = 'Times New Roman'  # 设置字体
            #     # run.font.size = Pt(11)  # 设置字号
            #     run.font.bold = 1  # 设置是否加粗
            #     # run.font.color.rgb = RGBColor.from_string('00BFFF')  # 设置文字颜色
            #     table.cell(r, c).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置居中
    # for r, row in enumerate(table.rows):
    #     for c in range(len(row.cells)):
    #         # print(c)
    #         row.cells[c].text = td[r][c]
    doc.save(path)

def doc2pdf(doc_name, pdf_name):
    import os
    """
    :word文件转pdf
    :param doc_name word文件名称
    :param pdf_name 转换后pdf文件名称
    """
    word = client.DispatchEx("Word.Application")
    try:
        if os.path.exists(pdf_name):
            os.remove(pdf_name)
        worddoc = word.Documents.Open(doc_name,ReadOnly = 1)
        worddoc.SaveAs(pdf_name, FileFormat = 17)
        worddoc.Close()
        return pdf_name
    except Exception as e:
        print(e)
    finally:
        word.Quit(constants.wdDoNotSaveChanges)



if __name__ == '__main__':
    import os
    basic_info_path = r'data/Demo_Qual Detect Mono Color/report_resources/basic_info.html'
    # td = read_html(basic_info_path)
    # table=get_basic_table(basic_info_path)
    # basic_table2doc(table)
    from win32com import client
    from win32com.client import Dispatch, constants, gencache

    doc_name = r"E:\part\part_time_job\yg\run_data.docx"

    pdf_name = r"E:\part\part_time_job\yg\all_data.pdf"
    doc2pdf(doc_name, pdf_name)
    # gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)

    # w = Dispatch("Word.Application")
    # if os.path.exists(pdf_name):
    #     os.remove(pdf_name)
    # try:
    #     worddoc = w.Documents.Open(doc_name, ReadOnly=1)
    #     worddoc.SaveAs(pdf_name, FileFormat=17)
    #     # worddoc.ExportAsFixedFormat(pdf_name, constants.wdExportFormatPDF,
    #     #                         Item=constants.wdExportDocumentWithMarkup,
    #     #                         CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    #
    #     worddoc.Close()
    # except Exception as e:
    #     print('err')
    #     print(e)
    # finally:
    #     w.Quit(constants.wdDoNotSaveChanges)
